# -*- coding: utf-8 -*-
"""Kiem tra bao cao da xuat co dung yeu cau hinh thuc cua mau khong."""
import zipfile
from docx import Document
from docx.shared import Cm, Pt
from docx.oxml.ns import qn

PATH = "BAO_CAO_TONG_KET_NCKH_Goldtrading.docx"
doc = Document(PATH)
ok, fail = [], []


def check(cond, desc, actual=""):
    (ok if cond else fail).append(f"{desc}" + (f"  [{actual}]" if actual else ""))


# ---------------- 1. Hinh thuc trang ----------------
print("=" * 74)
print("1. KIEM TRA HINH THUC TRANG (muc 2.1, 2.2 cua mau)")
print("=" * 74)
for i, s in enumerate(doc.sections):
    w_cm = round(s.page_width.cm, 1)
    h_cm = round(s.page_height.cm, 1)
    check(w_cm == 21.0 and h_cm == 29.7, f"Section {i}: kho giay A4",
          f"{w_cm} x {h_cm} cm")
    check(round(s.left_margin.cm, 1) == 3.0, f"Section {i}: le trai 3cm",
          f"{round(s.left_margin.cm,2)} cm")
    check(round(s.right_margin.cm, 1) == 2.0, f"Section {i}: le phai 2cm",
          f"{round(s.right_margin.cm,2)} cm")
    check(round(s.top_margin.cm, 1) == 2.0, f"Section {i}: le tren 2cm",
          f"{round(s.top_margin.cm,2)} cm")
    check(round(s.bottom_margin.cm, 1) == 2.0, f"Section {i}: le duoi 2cm",
          f"{round(s.bottom_margin.cm,2)} cm")

# ---------------- 2. Font va gian dong ----------------
print()
print("=" * 74)
print("2. KIEM TRA FONT VA GIAN DONG")
print("=" * 74)
n = doc.styles["Normal"]
check(n.font.name == "Times New Roman", "Font Normal = Times New Roman",
      str(n.font.name))
check(n.font.size == Pt(13), "Co chu Normal = 13", str(n.font.size.pt))
ls = n.paragraph_format.line_spacing
check(1.3 <= ls <= 1.5, "Gian dong trong khoang 1,3 - 1,5", str(ls))

# ---------------- 3. So trang phia tren, chinh giua ----------------
print()
print("=" * 74)
print("3. KIEM TRA SO TRANG (muc 2.3: chinh giua, phia tren)")
print("=" * 74)
hdr_with_page = 0
for i, s in enumerate(doc.sections):
    xml = s.header.paragraphs[0]._p.xml if s.header.paragraphs else ""
    has_page = "PAGE" in xml
    centered = 'w:val="center"' in xml
    if has_page:
        hdr_with_page += 1
        check(centered, f"Section {i}: so trang canh giua (header)")
check(hdr_with_page >= 2, "So trang dat trong HEADER (phia tren trang)",
      f"{hdr_with_page} section co PAGE field")

# ---------------- 4. Khong gach duoi ----------------
print()
print("=" * 74)
print("4. KIEM TRA KHONG GACH DUOI (muc 2.5)")
print("=" * 74)
underlined = []
for p in doc.paragraphs:
    for r in p.runs:
        if r.font.underline:
            underlined.append(p.text[:50])
for t in doc.tables:
    for row in t.rows:
        for c in row.cells:
            for p in c.paragraphs:
                for r in p.runs:
                    if r.font.underline:
                        underlined.append(p.text[:40])
check(not underlined, "Khong co van ban bi gach duoi",
      f"{len(underlined)} truong hop" if underlined else "0 truong hop")

# ---------------- 5. Cau truc theo muc 3 cua mau ----------------
print()
print("=" * 74)
print("5. KIEM TRA CAU TRUC BAO CAO (muc 3.1 - 3.13)")
print("=" * 74)
heads = [p.text.strip() for p in doc.paragraphs
         if p.style.name.startswith("Heading")]
required = [
    ("MỤC LỤC", "3.2 Muc luc"),
    ("DANH MỤC BẢNG BIỂU", "3.3 Danh muc bang bieu"),
    ("DANH MỤC HÌNH VẼ", "3.3 Danh muc hinh"),
    ("DANH MỤC NHỮNG TỪ VIẾT TẮT", "3.4 Danh muc tu viet tat"),
    ("MỞ ĐẦU", "3.5 Mo dau"),
    ("TỔNG QUAN TÌNH HÌNH NGHIÊN CỨU THUỘC LĨNH VỰC ĐỀ TÀI", "3.6 Tong quan"),
    ("LÝ DO LỰA CHỌN ĐỀ TÀI", "3.7 Ly do chon de tai"),
    ("MỤC TIÊU, NỘI DUNG, PHƯƠNG PHÁP NGHIÊN CỨU", "3.8 Muc tieu/ND/PP"),
    ("ĐỐI TƯỢNG VÀ PHẠM VI NGHIÊN CỨU", "3.9 Doi tuong & pham vi"),
    ("CHƯƠNG 1. CƠ SỞ LÝ THUYẾT", "3.10 Chuong 1"),
    ("CHƯƠNG 2. DỮ LIỆU VÀ PHƯƠNG PHÁP NGHIÊN CỨU", "3.10 Chuong 2"),
    ("CHƯƠNG 3. KẾT QUẢ NGHIÊN CỨU", "3.10 Chuong 3"),
    ("CHƯƠNG 4. THẢO LUẬN", "3.10 Chuong 4"),
    ("KẾT LUẬN VÀ KIẾN NGHỊ", "3.11 Ket luan & kien nghi"),
    ("TÀI LIỆU THAM KHẢO", "3.12 Tai lieu tham khao"),
    ("PHỤ LỤC", "3.13 Phu luc"),
]
for text, label in required:
    check(text in heads, f"{label}: '{text}'")

# Kiem tra Ket luan co ca phan A va B
check("A. KẾT LUẬN" in heads, "3.11a Phan ket luan")
check("B. KIẾN NGHỊ" in heads, "3.11b Phan kien nghi")

# ---------------- 6. Bang, hinh, field ----------------
print()
print("=" * 74)
print("6. KIEM TRA BANG BIEU, HINH VE VA CAC FIELD")
print("=" * 74)
xml_all = doc.element.xml
check(len(doc.tables) >= 20, "So bang trong bao cao", f"{len(doc.tables)} bang")

with zipfile.ZipFile(PATH) as z:
    imgs = [n for n in z.namelist() if n.startswith("word/media/")]
check(len(imgs) >= 8, "So hinh anh nhung trong file", f"{len(imgs)} hinh")

seq_bang = xml_all.count("SEQ Bảng")
seq_hinh = xml_all.count("SEQ Hình")
check(seq_bang >= 20, "Field SEQ danh so bang tu dong", f"{seq_bang} field")
check(seq_hinh >= 8, "Field SEQ danh so hinh tu dong", f"{seq_hinh} field")

check(xml_all.count("TOC") >= 3, "Field TOC (muc luc + 2 danh muc)",
      f"{xml_all.count('TOC')} field")

with zipfile.ZipFile(PATH) as z:
    settings = z.read("word/settings.xml").decode("utf-8")
check("updateFields" in settings, "Bat co updateFields de Word tu cap nhat")

# Kiem tra danh so trang La Ma cho phan danh muc
check('w:fmt="lowerRoman"' in xml_all, "Phan danh muc danh so i, ii, iii")
check('w:fmt="decimal"' in xml_all, "Phan noi dung danh so 1, 2, 3")

# ---------------- 7. Do dai ----------------
print()
print("=" * 74)
print("7. THONG KE NOI DUNG")
print("=" * 74)
words = sum(len(p.text.split()) for p in doc.paragraphs)
for t in doc.tables:
    for row in t.rows:
        for c in row.cells:
            words += len(c.text.split())
paras = len([p for p in doc.paragraphs if p.text.strip()])
print(f"  So doan van co noi dung : {paras}")
print(f"  Tong so tu (uoc tinh)   : {words:,}")
print(f"  So bang                 : {len(doc.tables)}")
print(f"  So hinh                 : {len(imgs)}")
print(f"  So heading              : {len(heads)}")
print(f"  Uoc so trang            : ~{words // 380 + len(imgs) + 6} trang")

# ---------------- Ket qua ----------------
print()
print("=" * 74)
print(f"KET QUA: {len(ok)} DAT / {len(fail)} CHUA DAT")
print("=" * 74)
if fail:
    print("\nCAC MUC CHUA DAT:")
    for f in fail:
        print(f"  [ ] {f}")
else:
    print("\nToan bo cac muc kiem tra deu DAT.")
print()
print("Chi tiet cac muc DAT:")
for o in ok:
    print(f"  [x] {o}")
