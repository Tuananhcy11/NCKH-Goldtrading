# -*- coding: utf-8 -*-
"""
Engine dung de xuat ban bao cao NCKH theo Mau 2021-SVNCKH-04-BCTK.

Yeu cau hinh thuc (tu template):
  - Kho giay A4 (210 x 297 mm)
  - Phong Times New Roman, co chu 13
  - Paragraph 1,3 - 1,5 line  -> dung 1,4
  - Le trai 3 cm; le tren, duoi, phai 2 cm
  - So trang o chinh giua trang, PHIA TREN
  - Khong gach duoi cac tu, cau
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT = "Times New Roman"
SIZE = Pt(13)
LINE = 1.4

# ---------------------------------------------------------------- helpers XML
def _el(tag, **attrs):
    e = OxmlElement(tag)
    for k, v in attrs.items():
        e.set(qn(k), v)
    return e


def add_field(paragraph, instr, placeholder="1"):
    """Chen mot field Word (PAGE, TOC, SEQ...) dang fldSimple."""
    fld = _el("w:fldSimple", **{"w:instr": instr})
    fld.set(qn("w:dirty"), "true")
    r = _el("w:r")
    rpr = _el("w:rPr")
    rf = _el("w:rFonts", **{"w:ascii": FONT, "w:hAnsi": FONT})
    sz = _el("w:sz", **{"w:val": "26"})
    rpr.append(rf)
    rpr.append(sz)
    r.append(rpr)
    t = _el("w:t")
    t.text = placeholder
    r.append(t)
    fld.append(r)
    paragraph._p.append(fld)
    return paragraph


def enable_update_fields(doc):
    """Bat co updateFields de Word tu cap nhat MUC LUC khi mo file."""
    settings = doc.settings.element
    for tag in ("w:updateFields",):
        existing = settings.find(qn(tag))
        if existing is not None:
            settings.remove(existing)
    settings.append(_el("w:updateFields", **{"w:val": "true"}))


# ---------------------------------------------------------------- document
def new_document():
    doc = Document()

    # --- Base style ---
    st = doc.styles["Normal"]
    st.font.name = FONT
    st.font.size = SIZE
    st.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    pf = st.paragraph_format
    pf.line_spacing = LINE
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)

    # --- Heading styles: Times New Roman, den, khong gach duoi ---
    for name, size, bold, italic in (
        ("Heading 1", 14, True, False),
        ("Heading 2", 13, True, False),
        ("Heading 3", 13, False, True),
        ("Caption", 12, False, True),
    ):
        s = doc.styles[name]
        s.font.name = FONT
        s.font.size = Pt(size)
        s.font.bold = bold
        s.font.italic = italic
        s.font.underline = False
        s.font.color.rgb = RGBColor(0, 0, 0)
        s.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        s.paragraph_format.line_spacing = LINE
        s.paragraph_format.space_before = Pt(10 if name.startswith("Heading") else 4)
        s.paragraph_format.space_after = Pt(6)
        s.paragraph_format.keep_with_next = True

    # --- Page setup A4 + le theo quy dinh ---
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.left_margin = Cm(3.0)
    sec.right_margin = Cm(2.0)
    sec.top_margin = Cm(2.0)
    sec.bottom_margin = Cm(2.0)
    sec.header_distance = Cm(1.2)
    sec.footer_distance = Cm(1.2)

    return doc


def set_page_number_header(section):
    """So trang o chinh giua, PHIA TREN trang."""
    hdr = section.header
    hdr.is_linked_to_previous = False
    p = hdr.paragraphs[0]
    p.text = ""
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    add_field(p, " PAGE ")
    for r in p.runs:
        r.font.name = FONT
        r.font.size = SIZE
    return p


def clear_header(section):
    hdr = section.header
    hdr.is_linked_to_previous = False
    hdr.paragraphs[0].text = ""


def new_section(doc, page_numbers=True, restart_at=None, fmt=None):
    """Tao section moi (dung de tach phan bia / phan danh muc / phan noi dung)."""
    sec = doc.add_section(WD_SECTION.NEW_PAGE)
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.left_margin = Cm(3.0)
    sec.right_margin = Cm(2.0)
    sec.top_margin = Cm(2.0)
    sec.bottom_margin = Cm(2.0)
    sec.header_distance = Cm(1.2)
    sec.footer_distance = Cm(1.2)

    if page_numbers:
        set_page_number_header(sec)
    else:
        clear_header(sec)

    if restart_at is not None or fmt is not None:
        pgnum = _el("w:pgNumType")
        if fmt:
            pgnum.set(qn("w:fmt"), fmt)
        if restart_at is not None:
            pgnum.set(qn("w:start"), str(restart_at))
        sec._sectPr.append(pgnum)
    return sec


# ---------------------------------------------------------------- content API
def para(doc, text, align="justify", bold=False, italic=False, size=None,
         first_line=1.0, space_after=6, space_before=0, line=None):
    p = doc.add_paragraph()
    p.alignment = {
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }[align]
    pf = p.paragraph_format
    pf.line_spacing = line or LINE
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    if first_line:
        pf.first_line_indent = Cm(first_line)
    r = p.add_run(text)
    r.font.name = FONT
    r.font.size = size or SIZE
    r.bold = bold
    r.italic = italic
    r.font.underline = False
    return p


def heading(doc, text, level=1, align="left", page_break=False):
    if page_break:
        doc.add_page_break()
    h = doc.add_heading(text, level=level)
    h.alignment = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }[align]
    for r in h.runs:
        r.font.name = FONT
        r.font.underline = False
        r.font.color.rgb = RGBColor(0, 0, 0)
    return h


def bullet(doc, text, level=0, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    pf = p.paragraph_format
    pf.line_spacing = LINE
    pf.space_after = Pt(4)
    pf.left_indent = Cm(1.0 + 0.6 * level)
    pf.first_line_indent = Cm(-0.4)
    if bold_prefix:
        rb = p.add_run(bold_prefix)
        rb.font.name = FONT
        rb.font.size = SIZE
        rb.bold = True
    r = p.add_run(text)
    r.font.name = FONT
    r.font.size = SIZE
    return p


def caption(doc, kind, title, above=True):
    """Chen caption dang 'Bang 1: ...' / 'Hinh 1: ...' voi field SEQ
    de Word tu dong danh so va sinh DANH MUC."""
    p = doc.add_paragraph(style="Caption")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.line_spacing = 1.0
    pf.space_before = Pt(6 if above else 4)
    pf.space_after = Pt(4 if above else 8)
    pf.keep_with_next = above

    r = p.add_run(f"{kind} ")
    r.font.name = FONT
    r.font.size = Pt(12)
    r.bold = True
    r.italic = False

    add_field(p, f" SEQ {kind} \\* ARABIC ")

    r2 = p.add_run(f": {title}")
    r2.font.name = FONT
    r2.font.size = Pt(12)
    r2.bold = False
    r2.italic = False
    return p


def table(doc, headers, rows, widths=None, caption_title=None,
          note=None, font_size=11, align_first_left=True):
    """Bang co vien, header in dam, tu dong danh so bang field SEQ."""
    if caption_title:
        caption(doc, "Bảng", caption_title, above=True)

    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False

    total = Cm(15.8)
    if widths is None:
        widths = [total / len(headers)] * len(headers)
    else:
        s = sum(widths)
        widths = [Cm(15.8 * w / s) for w in widths]

    hdr = t.rows[0]
    for i, (h, w) in enumerate(zip(headers, widths)):
        c = hdr.cells[i]
        c.width = w
        c.text = ""
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(2)
        r = p.add_run(h)
        r.font.name = FONT
        r.font.size = Pt(font_size)
        r.bold = True
        _shade(c, "D9E2F3")

    for row in rows:
        cells = t.add_row().cells
        for i, (val, w) in enumerate(zip(row, widths)):
            c = cells[i]
            c.width = w
            c.text = ""
            p = c.paragraphs[0]
            first = (i == 0 and align_first_left)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if first else WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)
            txt = str(val)
            bold = txt.startswith("**") and txt.endswith("**")
            if bold:
                txt = txt[2:-2]
            r = p.add_run(txt)
            r.font.name = FONT
            r.font.size = Pt(font_size)
            r.bold = bold

    if note:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(8)
        r = p.add_run(note)
        r.font.name = FONT
        r.font.size = Pt(11)
        r.italic = True
    return t


def _shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = _el("w:shd", **{"w:val": "clear", "w:color": "auto", "w:fill": hexcolor})
    tcPr.append(shd)


def figure(doc, path, caption_title, width_cm=15.0, source=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    r = p.add_run()
    r.add_picture(path, width=Cm(width_cm))
    caption(doc, "Hình", caption_title, above=False)
    if source:
        sp = doc.add_paragraph()
        sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sp.paragraph_format.line_spacing = 1.0
        sp.paragraph_format.space_after = Pt(10)
        sr = sp.add_run(source)
        sr.font.name = FONT
        sr.font.size = Pt(11)
        sr.italic = True
    return p


def toc(doc, instr):
    """Chen field TOC. Word se tu dong sinh noi dung khi mo file."""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = LINE
    add_field(p, instr, placeholder="[Nhấn Ctrl+A rồi F9 trong Word để cập nhật mục lục]")
    return p


def formula(doc, text, note=None):
    """Cong thuc: canh giua, in nghieng."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6 if not note else 2)
    r = p.add_run(text)
    r.font.name = FONT
    r.font.size = SIZE
    r.italic = True
    if note:
        np_ = doc.add_paragraph()
        np_.alignment = WD_ALIGN_PARAGRAPH.CENTER
        np_.paragraph_format.line_spacing = 1.0
        np_.paragraph_format.space_after = Pt(8)
        nr = np_.add_run(note)
        nr.font.name = FONT
        nr.font.size = Pt(11)
        nr.italic = True
    return p
